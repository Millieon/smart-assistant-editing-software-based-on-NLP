# encoding=utf-8
from docx import Document
file = Document("E:\PyCharm\EditorAssistance\TEST1.docx")     #读取文档内容
print(file.paragraphs[4].text)                    #输出文章标题
m=0;
for s in file.paragraphs[4].text:

  if s==';'or s=='；':
        m+=1;
m+=1;
print("关键词个数："+str(m))
