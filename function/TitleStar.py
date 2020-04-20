# encoding=utf-8


from docx import Document

file = Document("E:\PyCharm\EditorAssistance\TEST1.docx")  # 读取文档内容
print("标题是：" + file.paragraphs[0].text)  # 输出文章标题

m = 0;
for s in file.paragraphs[0].text:
    # 中文字符范围
    if s == '*':
        m += 1;


print("标题*个数为：" + str(m))  # 统计并输出标题中*的数量